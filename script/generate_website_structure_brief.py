from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Golden_Forests_Website_Structure_Brief_July_2026.docx"
GREEN = "17392E"
OLIVE = "6B8E23"
CREAM = "F4E8D2"
DARK = "1B1B1B"
MUTED = "5F685F"


def shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    node = properties.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        properties.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell, top=90, start=100, bottom=90, end=100):
    properties = cell._tc.get_or_add_tcPr()
    container = properties.first_child_found_in("w:tcMar")
    if container is None:
        container = OxmlElement("w:tcMar")
        properties.append(container)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = container.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            container.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def width(cell, inches):
    properties = cell._tc.get_or_add_tcPr()
    node = properties.find(qn("w:tcW"))
    if node is None:
        node = OxmlElement("w:tcW")
        properties.append(node)
    node.set(qn("w:w"), str(int(inches * 1440)))
    node.set(qn("w:type"), "dxa")


def repeat_header(row):
    properties = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    properties.append(node)


def field(paragraph, code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])


def table(doc, headers, rows, widths, font_size=8.2):
    result = doc.add_table(rows=1, cols=len(headers))
    result.style = "Table Grid"
    result.alignment = WD_TABLE_ALIGNMENT.CENTER
    result.autofit = False
    repeat_header(result.rows[0])
    for index, label in enumerate(headers):
        cell = result.rows[0].cells[index]
        cell.text = label
        shade(cell, GREEN)
        margins(cell)
        width(cell, widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8.5)
    for row_index, values in enumerate(rows):
        cells = result.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            margins(cells[index])
            width(cells[index], widths[index])
            if row_index % 2:
                shade(cells[index], "F7F3EA")
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor.from_string(DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return result


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def callout(doc, title, body):
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.cell(0, 0)
    shade(cell, CREAM)
    margins(cell, 150, 180, 150, 180)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN)
    paragraph.add_run("\n" + body)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_header(doc, name, route, purpose, audience, outcome, sources):
    doc.add_heading(name, level=1)
    table(
        doc,
        ["Item", "Brief"],
        [
            ("Route", route),
            ("Primary purpose", purpose),
            ("Primary audience", audience),
            ("Desired visitor outcome", outcome),
            ("Principal July sources", sources),
        ],
        [1.6, 5.7],
        9,
    )


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in [
        ("Title", 34, GREEN),
        ("Subtitle", 15, OLIVE),
        ("Heading 1", 22, GREEN),
        ("Heading 2", 15, OLIVE),
        ("Heading 3", 11, GREEN),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "GOLDEN FORESTS  |  WEBSITE STRUCTURE BRIEF"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(OLIVE)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("July 2026 working brief  •  Page ")
    field(footer, "PAGE")
    footer.add_run(" of ")
    field(footer, "NUMPAGES")
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    # Cover
    brand = doc.add_paragraph()
    brand.paragraph_format.space_before = Pt(75)
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run("GOLDEN FORESTS")
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(OLIVE)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Website Structure &\nContent Rationale Brief")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("A page-by-page explanation of the live website and its July 2026 documentary sources")
    cover_rows = [
        ("Prepared", "20 July 2026"),
        ("Scope", "Public Golden Forests website"),
        ("Source folder", r"E:\GOLDEN FORESTS AI\NEW FILES"),
        ("Live application", "https://goldenforests-ai.onrender.com"),
        ("Content baseline", "Commit 415b62d and July 2026 source set"),
    ]
    cover = doc.add_table(rows=len(cover_rows), cols=2)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, (label, value) in enumerate(cover_rows):
        cover.cell(index, 0).text = label
        cover.cell(index, 1).text = value
        shade(cover.cell(index, 0), GREEN)
        shade(cover.cell(index, 1), "F7F3EA")
        margins(cover.cell(index, 0), 110, 130, 110, 130)
        margins(cover.cell(index, 1), 110, 130, 110, 130)
        for run in cover.cell(index, 0).paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        for cell in cover.rows[index].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    purpose = doc.add_paragraph()
    purpose.alignment = WD_ALIGN_PARAGRAPH.CENTER
    purpose.paragraph_format.space_before = Pt(30)
    run = purpose.add_run("Purpose of this document")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN)
    paragraph = doc.add_paragraph(
        "To give management, marketing, compliance, design and development teams one shared explanation of what is on the website, why each element exists, and which July 2026 documents support it."
    )
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contents
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Contents", level=1)
    for item in [
        "1. Executive summary",
        "2. Source documents and content-governance rules",
        "3. Website sitemap and global structure",
        "4. Home",
        "5. Investment",
        "6. Precision Farming",
        "7. Asset Management",
        "8. Golden Forests Group",
        "9. Contact",
        "10. Legal pages",
        "11. Cross-site calls to action and document handling",
        "12. Content-maintenance checklist",
        "Appendix A. Source-to-page matrix",
    ]:
        paragraph = doc.add_paragraph(item)
        paragraph.paragraph_format.left_indent = Inches(0.15)
        paragraph.paragraph_format.space_after = Pt(5)
    callout(
        doc,
        "Document convention",
        "“PDF p.” means the page number in the PDF file as opened by a reader. A printed page number inside a PDF may differ where the cover is counted separately.",
    )

    # Executive summary
    doc.add_page_break()
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "The website is a concise, enquiry-led presentation of Golden Forests for eligible professional and corporate investors. It explains the underlying agarwood and mango plantation strategy without turning the public website into an offering document. The site introduces the opportunity, operating platform, risk controls, reporting model and team, then directs interested visitors to contact the company or view controlled supporting material."
    )
    bullets(
        doc,
        [
            "The website presents a proposed fund-share model, not direct ownership of individual trees.",
            "USD 100,000 is described only as an indicative reference subscription; final eligibility and minimums depend on the relevant platform and definitive documents.",
            "Tree-equivalent references are used for economic allocation, accounting and reporting—not to transfer ownership of specific trees, land or planting blocks.",
            "Agarwood and mango are crop-specific exposures with different harvest horizons; exact returns and disputed share prices are deliberately excluded from public pages.",
            "Operational claims are framed as intended, proposed, modelled or subject to implementation where the draft PPM says arrangements are not final.",
            "The principal conversion action is an information request, not an online investment or subscription.",
        ],
    )
    callout(
        doc,
        "Why the website is intentionally conservative",
        "The July files contain internal inconsistencies. The public website follows the cautious position in the draft PPM and the fund-share explanation in the July FAQ, and avoids presenting a disputed number or draft structure as final.",
    )

    # Sources
    doc.add_heading("2. Source Documents and Content-Governance Rules", level=1)
    source_rows = [
        ("6. GF Company Overview-July 2026.pdf", "PDF pp. 1–2", "Corporate overview, proposed structure, differentiators, crop/share-class summary, 1:1 reforestation and risk notice.", "Primary high-level narrative."),
        ("6. GF Agarwood Exposé July 2026.pdf", "PDF pp. 1–23", "Agarwood market, species, FIO inoculation, lifecycle, technology, operations, illustrative economics, risk and disclaimers.", "Crop detail; figures remain illustrative."),
        ("6. GF Mango Exposé July 2026.pdf", "PDF pp. 1–21", "Sweet Elena, market context, orchard model, lifecycle, technology, economics, deductions, risk and disclaimers.", "Crop detail; disputed price excluded."),
        ("6. GF Professional Client Investment presentation - July 2026.pdf", "PDF pp. 2–35", "Investor summary, USD 100,000 illustrations, sub-fund structure, operations, technology, risk, sustainability, reporting and next steps.", "Investor journey and reporting."),
        ("6. GF-FAQ-July 2026.pdf", "PDF pp. 1–13", "Company, fund shareholding, fees, documentation, reporting, replacement stock, visits, transfers and distributions.", "Primary public FAQ source."),
        ("9. Private Placement Memorandum-Professional Investors-Draft-July 2026.pdf", "PDF pp. 1–31", "Proposed architecture, illustrative ticket, economics, governance, reporting, risks, readiness, pricing schedules and environmental appendix.", "Controls cautious/draft status and qualifications."),
    ]
    table(doc, ["Document", "Pages", "What it contributes", "How it is used"], source_rows, [2.1, 0.8, 2.8, 1.55], 7.8)
    doc.add_heading("2.1 Source hierarchy used by the website", level=2)
    bullets(
        doc,
        [
            "Legal status, ownership, eligibility and finality: draft PPM first, then the July FAQ.",
            "High-level company narrative: Company Overview, qualified by the draft PPM.",
            "Crop science, lifecycle and operational detail: crop-specific exposés.",
            "Investor-friendly sequencing and visibility concepts: Professional Client Investment presentation.",
            "Where sources disagree, use the narrower and more qualified claim—or omit the figure entirely.",
        ],
    )
    doc.add_heading("2.2 Material conflicts and the website treatment", level=2)
    conflicts = [
        ("Ownership", "Some exposé passages refer to client-owned trees; Company Overview, FAQ and PPM describe fund shares without direct tree ownership.", "Fund shares; no ownership of specific trees, blocks, land or plantation assets.", "Company Overview p.1; FAQ p.6; PPM pp.5, 11–14."),
        ("Mango share reference", "Company Overview shows USD 371.64; Mango Exposé and PPM use USD 437.08.", "No mango per-share price is shown publicly.", "Company Overview p.1; Mango Exposé p.7; PPM pp.12, 24–25."),
        ("Minimum subscription", "FAQ/presentation state USD 100,000; PPM says it is illustrative and the actual minimum is not final.", "USD 100,000 is an indicative reference subscription subject to definitive terms.", "FAQ p.11; Presentation pp.4, 24; PPM p.12."),
        ("Returns/IRR", "Different pages show different annualised outcomes; disclaimer language rejects reliance on projections.", "No exact return percentage or calculator; only illustrative timing and fee mechanics.", "Agarwood pp.11, 23; Mango pp.8, 11, 21; Presentation pp.14, 20, 22–23; PPM pp.5–8."),
        ("GPS reporting", "Some documents imply individual-tree GPS; the PPM frames reporting at fund, block and asset-pool level.", "Fund-level traceability; no specific trees assigned to individual shareholders.", "FAQ p.9; Presentation p.31; PPM pp.13, 17."),
        ("Insurance/replacement", "Exposés use guarantee language; the PPM warns insurance may be limited, costly or unavailable.", "Insurance and 20% surplus stock are controls that reduce, not eliminate, risk.", "Agarwood p.20; Mango p.18; FAQ pp.8–9; PPM pp.20–21."),
        ("Launch profile", "Older site copy used July 2026 and 50,000 trees; PPM uses a 31,000-tree Q4 2026 illustrative profile.", "Planned Q4 2026; 23,000 agarwood + 8,000 mango; subject to readiness.", "Presentation p.22; PPM pp.3–4, 9, 12, 26–28."),
    ]
    table(doc, ["Topic", "Conflict", "Website treatment", "Key sources"], conflicts, [0.95, 2.45, 2.45, 1.7], 7.5)

    # Sitemap
    doc.add_page_break()
    doc.add_heading("3. Website Sitemap and Global Structure", level=1)
    sitemap = [
        ("Home", "/ or /home", "First orientation and high-level investor proposition.", "Continue to Investment or request information."),
        ("Investment", "/investment", "Agarwood, mango, diversification, fees, FAQs and risk.", "Enquire or view the July FAQ."),
        ("Precision Farming", "/precision-farming", "Operations, technology, research, risk, traceability and lifecycles.", "Understand operating capability."),
        ("Asset Management", "/asset-management", "Optional plantation access and visual operations context.", "Register interest in a visit."),
        ("Golden Forests Group", "/golden-forests-group", "Company, capabilities, commitments, leadership and board.", "Build institutional credibility."),
        ("Contact", "/contact", "Qualified enquiries, contact data and resources.", "Submit the form or open documents."),
        ("Disclaimer", "/disclaimer", "Risk, draft/finality and public-site boundaries.", "Ensure informed use."),
        ("Privacy Policy", "/privacy-policy", "Personal-data handling and rights.", "Privacy transparency."),
        ("Cookie Policy", "/cookie-policy", "Cookies, analytics and choices.", "Tracking transparency."),
    ]
    table(doc, ["Page", "Route", "Role in the journey", "Expected outcome"], sitemap, [1.25, 1.35, 2.75, 2.1], 8)
    doc.add_heading("3.1 Global navigation and shell", level=2)
    global_rows = [
        ("Left sidebar", "Six public navigation destinations.", "Keeps the opportunity → operations → organisation → enquiry journey visible.", "Information architecture based on Company Overview p.1 and Presentation pp.2–35."),
        ("Philippine Operations", "External CADI operations portal.", "Separates investor/corporate content from detailed Philippine operations.", "Company Overview p.1; PPM pp.4, 9, 13, 26–28; Presentation pp.25–31."),
        ("Mobile header", "Brand and menu control.", "Maintains usability on small screens.", "UX decision; no July documentary source."),
        ("Footer risk line", "Risk and no-guarantee reminder with Disclaimer link.", "Places a compliance reminder on every page.", "Company Overview p.2; exposé disclaimers; PPM pp.1, 18–21."),
        ("Footer legal links", "Disclaimer, Privacy, Cookie and Contact.", "Persistent access to legal and contact information.", "Website legal requirement; risk content supported by July documents."),
    ]
    table(doc, ["Element", "What is inside", "Why it is there", "Source basis"], global_rows, [1.25, 1.9, 2.35, 2.05], 8)

    # Home
    doc.add_page_break()
    page_header(doc, "4. Home", "/ and /home", "Introduce the organisation and opportunity without presenting an offer or detailed economics.", "Eligible professional and corporate investors encountering Golden Forests for the first time.", "Understand the fund-share proposition and move to Investment or Contact.", "Company Overview pp.1–2; Presentation pp.2–8, 24–31; PPM pp.1–5, 11–17.")
    home = [
        ("Hero", "Structured access, eligible investors, proposed fund shares, no direct tree ownership and Request Information.", "Corrects the former ownership model and sets audience, structure and action.", "Company Overview p.1; FAQ pp.6, 11; PPM pp.2–5, 11–14."),
        ("Two crop pathways", "Agarwood and Sweet Elena cards with tree-equivalent shares and harvest horizons; no disputed returns.", "Shows two crop sleeves with different timelines while remaining concise.", "Agarwood pp.10–13; Mango pp.7–10; Presentation pp.9–24; PPM pp.5–8, 12."),
        ("Value pillars", "Structured Fund Access, Professional Management and Environmental Impact.", "Summarises governance, capability and stewardship.", "Company Overview p.1; Presentation pp.24–31; PPM pp.11–17, 26–31."),
        ("Research credibility", "PRMSU, VSU and UPLB cards.", "Links technical claims to recognised agricultural research relationships.", "Company Overview p.1; Agarwood p.15; Mango pp.13–14."),
        ("Closing CTA", "Long-term exposure, disciplined operations, stewardship and Contact Our Team.", "Provides a non-transactional next step.", "Company Overview p.1; Presentation pp.33–35; PPM pp.16, 27–29."),
    ]
    table(doc, ["Section", "What is inside", "Why it is there", "Document source"], home, [1.15, 2.55, 2.25, 1.6], 7.9)
    callout(doc, "Deliberately absent", "No per-share prices, IRRs, guaranteed returns or direct subscribe button. Those details are disputed or subject to final offering terms.")

    # Investment
    doc.add_page_break()
    page_header(doc, "5. Investment", "/investment", "Give a qualified overview of both crop sleeves, diversification, FAQs and risk.", "Professional/corporate investors seeking more detail.", "Understand the opportunity, read the FAQ and submit an enquiry.", "All six July documents; especially crop exposés, FAQ pp.6–12 and PPM pp.5–25.")
    investment = [
        ("Hero", "Proposed fund shareholding, professional management and long-term participation.", "Replaces direct ownership / attractive returns framing.", "Company Overview p.1; FAQ p.6; PPM pp.1–5, 11–14."),
        ("Agarwood card", "Crop context, demand, years 9–10 model, 10% revenue share, inoculation and strengths.", "Explains commercial logic and current mechanics without promising results.", "Agarwood pp.2–15, 20–23; Presentation pp.9–15, 23–24; PPM pp.12, 14–16, 22–23."),
        ("Mango card", "Sweet Elena, premium market, years 5–25 model, 20% commission and deductions.", "Explains recurring-harvest exposure without the conflicting share price.", "Mango pp.2–13, 18–21; Presentation pp.16–24; PPM pp.12, 14–16, 24–25."),
        ("Diversification", "Different crop markets and timing; diversification does not remove risk.", "Explains why both sleeves may be considered together without guaranteeing benefit.", "Agarwood p.21; Mango p.19; Presentation p.22; PPM pp.5–8, 12."),
        ("Five FAQs", "Eligibility/minimum, ownership, documents, fees/distributions and losses/replacement.", "Answers the most material misunderstandings before enquiry.", "FAQ pp.6–12; PPM pp.11–17, 20–25."),
        ("View FAQ", "Opens the public July FAQ PDF in a new tab.", "Makes the complete approved FAQ immediately available.", "GF-FAQ-July 2026.pdf, pp.1–13."),
        ("Important Notice", "No offer/advice; proposed/illustrative; possible loss; illiquid; professional investors; definitive documents.", "Places the compliance boundary on the highest-risk page.", "Company Overview p.2; Agarwood p.23; Mango p.21; PPM pp.1, 18–21."),
        ("Closing CTA", "Operations and reporting statement linked to Precision Farming.", "Moves visitors from financial overview to operational diligence.", "Presentation pp.25–31; PPM pp.13, 17, 26–29."),
    ]
    table(doc, ["Section", "What is inside", "Why it is there", "Document source"], investment, [1.15, 2.6, 2.2, 1.6], 7.6)
    callout(doc, "Removed from the prior website", "The outdated calculator and two-pager, USD 5,887.50 minimum, old tree prices and 18.5% / 14–23% return claims were removed because they did not match the July fund documents or conflicted across them.")

    # Precision Farming
    doc.add_page_break()
    page_header(doc, "6. Precision Farming", "/precision-farming", "Explain how the plantations are intended to be operated, monitored, protected and reported.", "Investors and advisers performing operational due diligence.", "Understand operating capability and continue to the operations portal or Contact.", "Agarwood pp.9, 12–16, 20; Mango pp.5–6, 9, 12–18; Presentation pp.25–32; FAQ pp.7–9; PPM pp.13, 17, 26–29.")
    precision = [
        ("Operations overview", "Zambales model, security, irrigation and field protocols.", "Shows a real operating platform rather than only financial structuring.", "Company Overview p.1; Presentation pp.25–26; PPM pp.4, 9, 13, 26–28."),
        ("Agroforestry Intelligence", "Monitoring, analytics, irrigation, drones, sensors and targeted inputs.", "Explains “AI” as practical operating discipline.", "Agarwood p.14; Mango p.12; FAQ pp.7, 9; Presentation p.26."),
        ("University partners", "PRMSU, VSU and UPLB roles.", "Provides technical validation and research context.", "Company Overview p.1; Agarwood p.15; Mango pp.13–14."),
        ("Risk controls", "Insurance terms, 20% surplus stock, multiple sites and inspections.", "Shows risk management without promising complete protection.", "Agarwood p.20; Mango p.18; FAQ pp.8–9; Presentation p.27; PPM pp.20–21."),
        ("Environment", "1:1 native-tree intention, biodiversity, carbon and certification qualifications.", "Explains stewardship without presenting preliminary carbon as certified value.", "Company Overview p.1; Agarwood p.16; Mango p.15; Presentation pp.28–30; PPM pp.26–31."),
        ("Transparency", "Fund-level GPS, contemplated portal, quarterly updates, annual reporting and optional visits.", "Corrects the former “your tree / coordinates” model.", "FAQ pp.6, 9; Presentation p.31; PPM pp.13, 17."),
        ("Q4 2026 milestone", "Planned 31,000 trees: 23,000 agarwood + 8,000 mango; subject to completion.", "Provides a current planning reference, not a completion claim.", "Presentation p.22; PPM pp.3–4, 9, 12, 26–28."),
        ("Crop lifecycles", "Agarwood establishment/inoculation/years 9–10; mango establishment/year 5/year 10–25.", "Explains the different horizons and operating requirements.", "Agarwood pp.13–14; Mango pp.9–12; Presentation pp.12–21; PPM pp.12, 14–16."),
    ]
    table(doc, ["Section", "What is inside", "Why it is there", "Document source"], precision, [1.15, 2.6, 2.2, 1.6], 7.6)

    # Asset management
    doc.add_page_break()
    page_header(doc, "7. Asset Management", "/asset-management", "Explain optional physical access and visual/location context.", "Qualified investors seeking first-hand operational familiarisation.", "Request a visit or review operations.", "FAQ pp.6, 11; Presentation p.31; PPM p.16 and pp.26–28.")
    asset = [
        ("Hero", "Optional visits subject to operational availability.", "Sets expectations without promising a fixed itinerary.", "FAQ p.6; Presentation p.31; PPM p.16."),
        ("Operations context", "Zambales overview and Philippine portal link.", "Connects visits to the actual operating platform.", "PPM pp.4, 9, 13, 26–28; Presentation pp.25–26."),
        ("Visit explanation", "Potential guided access, briefings and discussions; itinerary/hospitality confirmed individually.", "Resolves inconsistent two-day, luxury hotel and travel-cost source language.", "FAQ pp.6, 11; Presentation p.31; PPM p.16."),
        ("Due-diligence qualifier", "A visit does not replace legal, financial, tax or operational advice.", "Prevents a visit from being treated as investment validation.", "PPM pp.1, 16, 18–21."),
        ("Gallery", "Zambales, accommodation, surveillance and Clark context.", "Makes location and operating environment tangible.", "PPM pp.4, 9, 26–28; imagery is a website asset."),
    ]
    table(doc, ["Section", "What is inside", "Why it is there", "Document source"], asset, [1.15, 2.6, 2.2, 1.6], 8)

    # About
    doc.add_page_break()
    page_header(doc, "8. Golden Forests Group", "/golden-forests-group", "Explain the organisation, differentiators, commitments and people.", "Professional investors, advisers and partners evaluating credibility.", "Understand the platform and responsible parties.", "Company Overview pp.1–2; FAQ pp.2–5; crop exposés; PPM pp.9–14, 26–29.")
    about = [
        ("Hero/overview", "Proposed Singapore VCC, crop-specific funds and no direct ownership.", "Defines company and structure before values or credentials.", "Company Overview p.1; FAQ p.6; PPM pp.2–5, 11–14."),
        ("Differentiators", "Experience, inoculation, DENR, universities, AI and reforestation.", "Turns key capabilities into scannable proof points.", "Company Overview p.1; Agarwood pp.8, 12, 15–16; Mango pp.4, 6, 13–15."),
        ("For investors", "Fund shares, reporting and illustrative harvest economics.", "Replaces direct ownership and fixed returns.", "Company Overview pp.1–2; FAQ pp.6, 9–12; PPM pp.11–17, 20–25."),
        ("For the land", "Native-tree intention and qualified environment metrics.", "Explains environmental purpose without overstatement.", "Company Overview p.1; Presentation pp.28–30; PPM pp.26–31."),
        ("For the people", "Employment, training, fair treatment and communities.", "Shows social outcomes alongside operations.", "FAQ pp.3–5; Company Overview p.1; Presentation pp.28–30."),
        ("Leadership", "Chairman, President/CEO and Marketing/Sales profiles.", "Shows strategy, operations and investor-relations responsibility.", "Approved website biographies; PPM p.10 provides management context."),
        ("Board", "Agroforestry, agri-science and commercial directors.", "Shows functional oversight across core disciplines.", "Approved biographies; disciplines align with Company Overview and exposés."),
    ]
    table(doc, ["Section", "What is inside", "Why it is there", "Document source"], about, [1.15, 2.6, 2.2, 1.6], 7.7)

    # Contact
    doc.add_page_break()
    page_header(doc, "9. Contact", "/contact", "Convert qualified interest into a controlled enquiry and provide current resources.", "Professional/corporate investors, advisers and partners.", "Submit an enquiry, contact offices or open documents.", "FAQ pp.11–12; Presentation pp.33–35; PPM pp.16, 18–19, 27–29; July exposés/FAQ.")
    contact = [
        ("Hero", "Discuss proposed fund shares and request current materials.", "Prevents the form from appearing to accept subscriptions.", "FAQ p.11; PPM pp.1–3, 16, 18–19."),
        ("Pipedrive form", "Third-party enquiry form.", "Creates a controlled investor-relations hand-off.", "Presentation pp.33–35; PPM pp.16, 27–29; CRM is an implementation choice."),
        ("Contact cards", "UAE office, Philippine office, phone and email.", "Provides corporate and operational transparency.", "Company Overview p.1 footprint; addresses are approved website data."),
        ("Investor resources", "Request Agarwood, Request Mango and View FAQ.", "Controls detailed material while making the public FAQ accessible.", "Agarwood pp.1–23; Mango pp.1–21; FAQ pp.1–13."),
        ("Public FAQ file", "GF-FAQ-July-2026.pdf in a new tab.", "Makes the complete July FAQ available while the page summary remains concise.", "6. GF-FAQ-July 2026.pdf."),
    ]
    table(doc, ["Section", "What is inside", "Why it is there", "Document source"], contact, [1.15, 2.6, 2.2, 1.6], 8)
    callout(doc, "Document-access rule", "The FAQ is public. Crop exposés remain request-gated. The draft PPM and Professional Client Investment presentation are not unrestricted public downloads because their eligibility, distribution and confidentiality conditions require control.")

    # Legal
    doc.add_page_break()
    doc.add_heading("10. Legal Pages", level=1)
    legal = [
        ("Risk Warning and Disclaimer", "/disclaimer", "No offer/advice; illustrative; proposed structure; definitive documents; fund shares; illiquidity; principal risks and jurisdictions.", "Defines how investment statements must be read.", "Company Overview p.2; Agarwood p.23; Mango p.21; PPM pp.1, 18–21."),
        ("Privacy Policy", "/privacy-policy", "Data collected, purposes, legal bases, processors, transfers, retention, rights and contact.", "Explains enquiry-data handling.", "Operational/legal requirement; July files support only the enquiry context."),
        ("Cookie Policy", "/cookie-policy", "Essential, analytics and functional cookies; GA4, Clarity, choices and retention.", "Explains tracking and controls.", "Operational/legal requirement; no direct July source."),
    ]
    table(doc, ["Page", "Route", "What is inside", "Why it is there", "Source basis"], legal, [1.3, 1.05, 2.35, 1.55, 1.55], 7.7)
    callout(doc, "Legal maintenance point", "The disclaimer is a website summary only. It never replaces final offering, subscription, constitutional or jurisdiction-specific private-placement documents.")

    # CTAs
    doc.add_heading("11. Cross-Site Calls to Action and Document Handling", level=1)
    ctas = [
        ("Request Information / Contact Our Team", "Home and closing panels", "/contact", "Starts a conversation rather than accepting an investment."),
        ("Enquiry and Information Requests", "Home/Investment crop cards", "/contact", "One consistent next step across both crops."),
        ("Explore Plantation Operations", "Operations pages", "External CADI portal", "Moves visitors to the Philippine platform."),
        ("View FAQ Document", "Investment and Contact", "/GF-FAQ-July-2026.pdf", "Immediate access to the complete public FAQ."),
        ("Request crop exposés", "Contact", "Pipedrive request forms", "Controls distribution and records the recipient."),
        ("Retired two-pager", "Legacy direct URL", "302 to /contact", "Prevents outdated material circulating as current."),
    ]
    table(doc, ["CTA", "Location", "Destination", "Reason"], ctas, [1.7, 1.55, 2.0, 2.6], 8)

    # Maintenance
    doc.add_page_break()
    doc.add_heading("12. Content-Maintenance Checklist", level=1)
    bullets(
        doc,
        [
            "Every ownership statement must say fund shares or tree-equivalent references—not direct ownership of individual trees.",
            "Do not publish a mango share price until the USD 371.64 versus USD 437.08 conflict is formally resolved.",
            "Describe USD 100,000 as indicative unless final offering terms establish an actual minimum.",
            "Do not restore an IRR or calculator without one approved model, consistent definitions and compliance sign-off.",
            "Keep proposed, intended, illustrative, modelled, subject to and if declared where arrangements or outcomes are not final.",
            "Describe GPS visibility at fund, block or asset-pool level unless definitive documents support investor-specific tree allocation.",
            "Describe insurance and replacement stock as controls, never as complete protection or guaranteed recovery.",
            "Update Q4 2026 / 31,000 trees only when a newer approved readiness document supersedes the July draft PPM.",
            "Keep the draft PPM and confidential presentation distribution-controlled.",
            "When replacing the FAQ PDF, retain its public URL or add a redirect.",
            "Update the Disclaimer whenever structure, eligibility, risk or governing documents change.",
            "Record the source file, version and PDF pages for every material website change.",
        ],
    )

    doc.add_heading("Appendix A. Source-to-Page Matrix", level=1)
    matrix = [
        ("Company Overview", "Home; Group; Investment; Precision Farming; Contact", "Corporate narrative, structure, differentiators, crops, reforestation, footprint and high-level risks."),
        ("Agarwood Exposé", "Home; Investment; Precision Farming; Contact", "Agarwood context, lifecycle, inoculation, technology, economics, risk and resource."),
        ("Mango Exposé", "Home; Investment; Precision Farming; Contact", "Sweet Elena, orchard model, lifecycle, technology, economics, deductions, risk and resource."),
        ("Professional presentation", "Home; Investment; Precision Farming; Asset Management; Contact", "Investor journey, ticket, sub-fund framing, reporting, technology, sustainability and next steps."),
        ("July FAQ", "Investment; Precision Farming; Asset Management; Group; Contact", "Ownership, fees, documentation, reporting, visits, replacement, transfer and distributions."),
        ("Draft PPM", "All investment-facing pages and Disclaimer", "Finality, architecture, illustrative status, launch, reporting, risks and document control."),
    ]
    table(doc, ["Source", "Website pages influenced", "Primary contribution"], matrix, [1.85, 2.65, 3.35], 8)
    doc.add_heading("Appendix B. Public and Non-Public Areas", level=1)
    bullets(
        doc,
        [
            "Public: all routes above, the July FAQ PDF and external Philippine operations link.",
            "Restricted: /admin/login and /admin, which manage website content and require credentials.",
            "Not unrestricted public downloads: draft PPM, Professional Client Investment presentation and crop exposés.",
        ],
    )
    ending = doc.add_paragraph()
    ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ending.paragraph_format.space_before = Pt(20)
    run = ending.add_run("End of brief")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(OLIVE)

    properties = doc.core_properties
    properties.title = "Golden Forests Website Structure & Content Rationale Brief"
    properties.subject = "Page-by-page website brief mapped to July 2026 source documents"
    properties.author = "Golden Forests / OpenAI Codex"
    properties.keywords = "Golden Forests, website structure, July 2026, source mapping"
    properties.comments = "Prepared from the live website and the six PDFs in NEW FILES."
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(path)
    print(path.stat().st_size)
