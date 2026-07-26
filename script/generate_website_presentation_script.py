from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Golden_Forests_Website_Presentation_Script_July_2026.docx"
GREEN = "17392E"
OLIVE = "6B8E23"
CREAM = "F4E8D2"
LIGHT_GREEN = "E8F0E3"
LIGHT_BLUE = "E8F0F5"
DARK = "1B1B1B"
MUTED = "5F685F"


def shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    node = properties.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        properties.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell, top=130, start=150, bottom=130, end=150):
    properties = cell._tc.get_or_add_tcPr()
    container = properties.first_child_found_in("w:tcMar")
    if container is None:
        container = OxmlElement("w:tcMar")
        properties.append(container)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = container.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            container.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def field(paragraph, code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])


def label_box(doc, label, text, fill=CREAM):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    margins(cell)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN)
    paragraph.add_run("\n" + text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def action(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(OLIVE)
    run.font.size = Pt(10)


def say(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(DARK)


def source_note(doc, text):
    label_box(doc, "SOURCE NOTE — DO NOT READ ALOUD", text, LIGHT_BLUE)


def transition(doc, text):
    label_box(doc, "TRANSITION", text, LIGHT_GREEN)


def qa(doc, question, answer):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, "F7F3EA")
    margins(cell)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("IF ASKED: " + question)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN)
    paragraph.add_run("\n" + answer)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def script_heading(doc, number, title, route, timing):
    doc.add_heading(f"{number}. {title}", level=1)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(f"ON SCREEN: {route}   |   Suggested time: {timing}")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(OLIVE)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [
        ("Title", 34, GREEN),
        ("Subtitle", 15, OLIVE),
        ("Heading 1", 22, GREEN),
        ("Heading 2", 15, OLIVE),
        ("Heading 3", 11, GREEN),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "GOLDEN FORESTS  |  WEBSITE PRESENTATION SCRIPT"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(OLIVE)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Presenter script  •  July 2026  •  Page ")
    field(footer, "PAGE")
    footer.add_run(" of ")
    field(footer, "NUMPAGES")
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    # Cover
    brand = doc.add_paragraph()
    brand.paragraph_format.space_before = Pt(80)
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run("GOLDEN FORESTS")
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(OLIVE)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Website Presentation Script")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("A natural, page-by-page script for presenting the July 2026 website")
    label_box(
        doc,
        "PRESENTATION FORMAT",
        "Suggested duration: 20–25 minutes\nAudience: professional investors, corporate investors, advisers and business partners\nLive website: https://goldenforests-ai.onrender.com\nContent baseline: commit 415b62d",
    )
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.paragraph_format.space_before = Pt(28)
    run = intro.add_run("This document is written to be spoken aloud.")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN)
    paragraph = doc.add_paragraph(
        "Green instructions tell you what to click or show. Main paragraphs are the words to say. Source notes and Q&A boxes are for your preparation and should not be read unless useful."
    )
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Run sheet
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Presentation Run Sheet", level=1)
    rows = [
        ("Opening", "1 minute"),
        ("Navigation and positioning", "1 minute"),
        ("Home", "3 minutes"),
        ("Investment", "5 minutes"),
        ("Precision Farming", "4 minutes"),
        ("Asset Management", "2 minutes"),
        ("Golden Forests Group", "3 minutes"),
        ("Contact and resources", "2 minutes"),
        ("Legal/footer and close", "2 minutes"),
        ("Questions", "As needed"),
    ]
    run_table = doc.add_table(rows=1, cols=2)
    run_table.style = "Table Grid"
    run_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, label in enumerate(["Section", "Suggested time"]):
        run_table.cell(0, index).text = label
        shade(run_table.cell(0, index), GREEN)
        for run in run_table.cell(0, index).paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for left, right in rows:
        cells = run_table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        margins(cells[0], 80, 120, 80, 120)
        margins(cells[1], 80, 120, 80, 120)
    doc.add_heading("How to use the script", level=2)
    for item in [
        "[CLICK] means select a menu item or button.",
        "[SCROLL] means move slowly enough for the audience to see the section before speaking.",
        "[PAUSE] is a deliberate pause for emphasis or questions.",
        "Do not quote a return percentage, mango share price or guaranteed outcome outside the wording in this script.",
        "If time is short, present Home, Investment, Precision Farming and Contact, then use the remaining pages only for questions.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    label_box(
        doc,
        "IMPORTANT PRESENTER RULE",
        "Say “fund shares,” “tree-equivalent references,” “indicative,” “illustrative,” “proposed” and “subject to definitive documents.” Do not say that an investor directly owns a tree or that a return, distribution, insurance claim or replacement is guaranteed.",
    )

    # Opening
    doc.add_page_break()
    script_heading(doc, "1", "Opening", "Home page loaded at the top", "1 minute")
    action(doc, "[ON SCREEN: HOME HERO. PAUSE FOR TWO SECONDS.] ")
    say(
        doc,
        "Good morning, and thank you for joining me. Today I am going to walk you through the Golden Forests website and explain how each page supports the overall investor journey—from understanding the opportunity, to reviewing plantation operations, to learning about the company and finally requesting further information.",
    )
    say(
        doc,
        "The website is designed for eligible professional and corporate investors. It is an information and due-diligence starting point; it is not an online offer, it does not accept subscriptions, and it does not replace the final legal and investment documents.",
    )
    say(
        doc,
        "The most important structural point is that the current model is based on fund shares. A share may use a tree-equivalent reference for allocation and reporting, but shareholders do not directly own a specific tree, planting block, land or plantation asset.",
    )
    source_note(doc, "Company Overview PDF p.1; July FAQ PDF pp.6 and 11; draft PPM PDF pp.1–5 and 11–14.")
    transition(doc, "I will begin with the navigation, because it shows the logic of the whole website.")

    # Navigation
    script_heading(doc, "2", "Navigation and Website Positioning", "Left sidebar and page header", "1 minute")
    action(doc, "[POINT TO EACH SIDEBAR ITEM WITHOUT CLICKING YET.] ")
    say(
        doc,
        "The navigation follows a simple sequence. Home introduces the proposition. Investment explains the two crop pathways. Precision Farming shows how the plantations are operated. Asset Management explains site access. Golden Forests Group introduces the organisation and its people. Contact is where a qualified visitor can make an enquiry or access current resources.",
    )
    say(
        doc,
        "At the bottom, the Philippine Operations link opens the external CADI operating portal. That separation is intentional: this website handles the corporate and investor narrative, while the external portal provides the detailed Philippine operational context.",
    )
    say(
        doc,
        "The footer carries a permanent risk reminder and links to the Disclaimer, Privacy Policy, Cookie Policy and Contact page, so those protections remain accessible throughout the journey.",
    )
    source_note(doc, "Company Overview PDF p.1; Professional Client Investment presentation PDF pp.25–31; draft PPM PDF pp.4, 9, 13 and 26–28.")
    transition(doc, "With that structure in mind, let us start with the Home page.")

    # Home
    script_heading(doc, "3", "Home Page", "/ or /home", "3 minutes")
    action(doc, "[CLICK HOME. RETURN TO THE TOP OF THE PAGE.] ")
    say(
        doc,
        "The Home page opens with the phrase ‘Structured access to sustainable agroforestry.’ This tells visitors immediately that Golden Forests is not presenting a simple tree purchase. The opportunity is structured for eligible professional and corporate investors through a proposed fund-share model.",
    )
    say(
        doc,
        "The Request Information button is deliberately the main action. We want a visitor to begin a controlled conversation with the team rather than make an investment decision from public website copy alone.",
    )
    action(doc, "[SCROLL TO THE TWO CROP PATHWAYS.] ")
    say(
        doc,
        "The next section introduces the two crop-specific pathways. Agarwood is the longer-cycle, concentrated realisation sleeve, with harvest activity modelled around years nine and ten. Sweet Elena Mango is the recurring-production sleeve, modelled to begin producing from year five and continue through year twenty-five. These are illustrative operating horizons, not promised distributions.",
    )
    say(
        doc,
        "Notice that the public cards do not show a mango share price or a headline return percentage. That is intentional. The July documents contain different mango reference prices and different performance illustrations, so the website avoids publishing a disputed number as though it were final.",
    )
    action(doc, "[SCROLL TO THE THREE VALUE PILLARS.] ")
    say(
        doc,
        "The three pillars summarise the proposition. First, Structured Fund Access: investor rights and obligations are governed by the final transaction documents. Second, Professional Management: the platform uses plantation teams, monitoring technology and scientific collaboration. Third, Environmental Impact: the operating model intends to plant one native Philippine tree for each corresponding underlying commercial tree represented in the fund allocation model.",
    )
    action(doc, "[SCROLL TO THE UNIVERSITY PARTNER CARDS.] ")
    say(
        doc,
        "The research partnership section supports the technical credibility of the plantation model. It highlights relationships with President Ramon Magsaysay State University, Visayas State University and the University of the Philippines Los Baños across mango research, soil science, pest management and propagation.",
    )
    action(doc, "[SCROLL TO THE FINAL HOME CTA.] ")
    say(
        doc,
        "The page closes with the central message: long-term agroforestry exposure, disciplined operations and responsible stewardship. From here, a visitor can contact the team or continue into the detailed Investment page.",
    )
    source_note(doc, "Company Overview PDF p.1; Agarwood Exposé PDF pp.10–16; Mango Exposé PDF pp.7–15; Professional presentation PDF pp.9–31; draft PPM PDF pp.5–17 and 26–31.")
    qa(doc, "Why do we not show returns on the Home page?", "Because the July documents contain different illustrations and the final economics remain subject to definitive documentation. The Home page explains the model and timing without presenting a projection as a promised outcome.")
    transition(doc, "Now that the high-level proposition is clear, I will open the Investment page and explain each crop in more detail.")

    # Investment
    script_heading(doc, "4", "Investment Page", "/investment", "5 minutes")
    action(doc, "[CLICK INVESTMENT. START AT THE HERO.] ")
    say(
        doc,
        "The Investment page is the main commercial overview. The hero now describes proposed fund shareholding, professional management and long-term agroforestry exposure. This replaces the former direct-tree-ownership message and keeps the public page aligned with the current fund documents.",
    )
    action(doc, "[SCROLL TO THE AGARWOOD PROGRAMME CARD.] ")
    say(
        doc,
        "The Agarwood section explains Aquilaria crassna and the international demand for agarwood and oud products. It then shows the operating model: establishment and cultivation, later inoculation and realisation modelled across years nine and ten.",
    )
    say(
        doc,
        "Under the current illustrative model, realised agarwood sales are subject to a ten percent Agarwood Management Revenue Share and applicable costs. The important point is that outcomes depend on actual yield, quality, market price, execution and the final governing documents. They are not guaranteed.",
    )
    say(
        doc,
        "The supporting strengths include licensed Forestry Industry Organization inoculation technology, Thai genetic sourcing, DENR permitting, CITES-related supply constraints and scientific operating relationships.",
    )
    action(doc, "[SCROLL TO THE MANGO PROGRAMME CARD.] ")
    say(
        doc,
        "The Mango section introduces the Sweet Elena Carabao variety, high-density orchard design and year-round production strategy. Production is modelled to begin from year five and continue through year twenty-five, subject to actual orchard performance.",
    )
    say(
        doc,
        "The current base case applies a twenty percent Mango Harvesting Commission to gross mango sales, together with permitted crop-specific deductions such as audit, reporting, logistics, certification, maintenance and harvesting costs. Any shareholder distribution would depend on realised proceeds, reserves, the governing documents and applicable law.",
    )
    action(doc, "[SCROLL TO THE DIVERSIFICATION SECTION.] ")
    say(
        doc,
        "The diversification section explains why the two crops may complement one another. Mango has an earlier recurring production profile, while agarwood has a later concentrated realisation profile. They also serve different markets. However, diversification can reduce concentration; it cannot eliminate biological, market, operational, liquidity or regulatory risk.",
    )
    action(doc, "[SCROLL TO FREQUENTLY ASKED QUESTIONS. OPEN EACH ITEM BRIEFLY.] ")
    say(
        doc,
        "The five displayed questions answer the most important points before a visitor contacts us. The opportunity is intended for eligible professional and corporate investors. The current material uses one hundred thousand US dollars as an indicative reference subscription, but the actual minimum and eligibility requirements are determined by the relevant platform and definitive offering documents.",
    )
    say(
        doc,
        "The ownership answer is equally important: the investor would hold shares in the relevant fund compartment. Tree-equivalent references support economic allocation and reporting, but do not give direct ownership of an individual tree or plantation asset.",
    )
    say(
        doc,
        "The remaining questions cover subscription and shareholder documentation, the fee and distribution structure, and how insurance and approximately twenty percent surplus replacement stock are intended to manage—rather than remove—plantation risk.",
    )
    action(doc, "[CLICK VIEW FAQ DOCUMENT. SHOW THAT THE JULY FAQ OPENS IN A NEW TAB, THEN RETURN.] ")
    say(
        doc,
        "The full July FAQ is available directly from the website. This gives visitors the complete set of plain-language answers while keeping the main Investment page concise.",
    )
    action(doc, "[SCROLL TO IMPORTANT NOTICE.] ")
    say(
        doc,
        "The Important Notice sets the legal boundary. The website is for information only; it is not an offer, solicitation or investment advice. The structure and all illustrative economics remain subject to legal, regulatory, tax, structuring and commercial review. The interests are long-term and illiquid, and investors may lose some or all of their capital.",
    )
    source_note(doc, "Agarwood Exposé PDF pp.2–15 and 20–23; Mango Exposé PDF pp.2–13 and 18–21; July FAQ PDF pp.6–12; Professional presentation PDF pp.9–24; draft PPM PDF pp.5–25.")
    qa(doc, "What is the minimum investment?", "The current materials use USD 100,000 as an indicative reference subscription. The actual minimum, allocation and eligibility are subject to the relevant platform and final offering terms.")
    qa(doc, "What return should I quote?", "Do not quote a fixed return from the public website. Explain the illustrative crop timelines and say that all outcomes depend on realised operations, prices, costs and definitive documents.")
    qa(doc, "Why is there no investment calculator?", "It was removed because the July files contain different pricing and performance assumptions. A public calculator should not be restored until one approved model is final and compliance-cleared.")
    transition(doc, "The Investment page explains the structure. The next page explains how the underlying plantation operations are intended to work.")

    # Precision Farming
    script_heading(doc, "5", "Precision Farming", "/precision-farming", "4 minutes")
    action(doc, "[CLICK PRECISION FARMING. START AT THE HERO.] ")
    say(
        doc,
        "Precision Farming is the operational due-diligence page. It explains how Golden Forests intends to manage agarwood and mango plantations in the Philippines through field teams, scientific collaboration and Agroforestry Intelligence.",
    )
    action(doc, "[SCROLL THROUGH ZAMBALES OPERATIONS AND AGROFORESTRY INTELLIGENCE.] ")
    say(
        doc,
        "The operations overview describes plantation preparation, irrigation, security and crop-specific protocols. Agroforestry Intelligence then explains the practical role of monitoring technology: sensors, drones, geotagging, targeted irrigation and fertilisation, pest detection, growth tracking and yield planning.",
    )
    action(doc, "[SCROLL TO UNIVERSITY PARTNERSHIPS.] ")
    say(
        doc,
        "The university section shows how research relationships support the operating model. PRMSU is associated with Sweet Elena and mango-development work. VSU supports soil science and integrated pest management. UPLB supports propagation and post-harvest or certification pathways.",
    )
    action(doc, "[SCROLL TO OPERATIONAL RISK CONTROLS.] ")
    say(
        doc,
        "The risk-control section has deliberately careful wording. Plantation insurance is subject to insurer availability, policy limits, exclusions and claims assessment. The model includes approximately twenty percent surplus planting stock for mortality replacement. Multiple sites and professional inspections are intended to reduce risk, but none of these controls can guarantee an investor outcome.",
    )
    action(doc, "[SCROLL TO ENVIRONMENTAL COMMITMENT.] ")
    say(
        doc,
        "The environmental section explains the intended one-to-one native-tree programme, biodiversity monitoring, carbon tracking and certification pathways. Preliminary carbon or biodiversity estimates are not presented as certified credits and are not included as guaranteed investor proceeds.",
    )
    action(doc, "[SCROLL TO INVESTOR TRANSPARENCY.] ")
    say(
        doc,
        "Investor Transparency replaces the old ‘your trees, your coordinates’ language. GPS and geotagged records support fund-managed inventory, planting blocks and biological-asset pools. The reporting framework contemplates secure portal visibility, quarterly operational updates, annual reporting where applicable and material-event notices.",
    )
    action(doc, "[SCROLL TO THE Q4 2026 MILESTONE AND LIFECYCLES.] ")
    say(
        doc,
        "The current illustrative launch profile targets fourth-quarter 2026 deployment of thirty-one thousand trees—twenty-three thousand agarwood and eight thousand mango—subject to legal, administrative and operational completion. The lifecycle diagrams then show why agarwood and mango have different operating and harvest horizons.",
    )
    source_note(doc, "Agarwood Exposé PDF pp.9 and 12–20; Mango Exposé PDF pp.5–6, 9 and 12–18; Professional presentation PDF pp.25–32; FAQ PDF pp.7–9; draft PPM PDF pp.13, 17 and 26–31.")
    qa(doc, "Does each shareholder receive GPS coordinates for a specific tree?", "No. The website describes fund-level, block-level and biological-asset-pool traceability. Tree references do not assign or transfer ownership of a specific tree to an investor.")
    qa(doc, "Is the 31,000-tree deployment complete?", "No. It is the current illustrative Q4 2026 launch profile and remains subject to legal, administrative and operational readiness.")
    transition(doc, "After explaining how the plantations are managed, the next page shows how qualified investors may obtain first-hand operational access.")

    # Asset Management
    script_heading(doc, "6", "Asset Management", "/asset-management", "2 minutes")
    action(doc, "[CLICK ASSET MANAGEMENT. SHOW THE HERO AND OPERATIONS CARD.] ")
    say(
        doc,
        "The Asset Management page explains the optional plantation-visit programme. Eligible professional and corporate investors may request access to operating sites and, where applicable, view assets associated with the relevant fund.",
    )
    say(
        doc,
        "The visit may include guided plantation access, operational briefings and discussions with management or technical teams. The itinerary, timing, access and hospitality are confirmed individually and remain subject to plantation conditions and operational availability.",
    )
    say(
        doc,
        "This wording is intentionally flexible because the July documents do not describe the visit package consistently. Most importantly, a plantation visit supports familiarisation and due diligence; it does not replace independent legal, financial, tax or operational advice.",
    )
    action(doc, "[SCROLL THROUGH THE FOUR GALLERY CARDS.] ")
    say(
        doc,
        "The gallery provides geographic and practical context: Southern Zambales, nearby accommodation, the plantation-monitoring environment and access through Clark. The page can then direct a visitor to Contact to register interest.",
    )
    source_note(doc, "July FAQ PDF pp.6 and 11; Professional presentation PDF p.31; draft PPM PDF p.16 and pp.26–28.")
    transition(doc, "We have now covered the opportunity and its operations. The next page introduces the organisation and the people behind it.")

    # Group
    script_heading(doc, "7", "Golden Forests Group", "/golden-forests-group", "3 minutes")
    action(doc, "[CLICK GOLDEN FORESTS GROUP. START AT THE HERO.] ")
    say(
        doc,
        "The Golden Forests Group page is the corporate credibility page. It defines Golden Forests as a platform providing professionally managed agarwood and mango exposure in the Philippines through a proposed Singapore Variable Capital Company structure with crop-specific fund compartments.",
    )
    say(
        doc,
        "It again states that shareholders would participate through fund shares rather than direct ownership of individual trees or plantation assets. Final rights, protections and economics are governed by the definitive offering, subscription and constitutional documents.",
    )
    action(doc, "[SCROLL THROUGH THE SIX DIFFERENTIATOR CARDS.] ")
    say(
        doc,
        "The differentiators bring together the main reasons for institutional interest: more than eighty years of combined management experience; access to licensed inoculation technology and Thai genetics; DENR permitting; relationships with Philippine universities; AI-enabled plantation intelligence; and the intended one-to-one native-tree programme.",
    )
    action(doc, "[SCROLL TO OUR COMMITMENT.] ")
    say(
        doc,
        "The commitment section is organised around investors, the land and the people. For investors, it highlights fund shareholding, reporting and illustrative harvest economics. For the land, it describes native-tree planting and environmental measurement. For people, it focuses on local employment, training, fair treatment and community value.",
    )
    action(doc, "[SCROLL THROUGH LEADERSHIP AND BOARD PROFILES.] ")
    say(
        doc,
        "The leadership and board sections show the people responsible for strategy, plantation operations, agricultural science, commercial development and investor relationships. These profiles are here to support governance and accountability, not simply brand presentation.",
    )
    source_note(doc, "Company Overview PDF pp.1–2; July FAQ PDF pp.2–6; Agarwood Exposé PDF pp.8, 12 and 15–16; Mango Exposé PDF pp.4, 6 and 13–15; draft PPM PDF pp.9–14 and 26–31.")
    transition(doc, "The final step in the visitor journey is Contact, where interest becomes a controlled enquiry.")

    # Contact
    script_heading(doc, "8", "Contact and Investor Resources", "/contact", "2 minutes")
    action(doc, "[CLICK CONTACT. SHOW THE HERO AND FORM.] ")
    say(
        doc,
        "The Contact page is the website’s main conversion point. The hero invites eligible visitors to discuss the proposed fund-share structure and request current professional-investor materials.",
    )
    say(
        doc,
        "The Pipedrive form creates a controlled hand-off to the company. It is an enquiry form, not an application or subscription form. This allows the team to understand the visitor, provide appropriate material and follow the relevant eligibility and compliance process.",
    )
    action(doc, "[SCROLL TO THE CONTACT CARDS.] ")
    say(
        doc,
        "The contact cards provide the UAE office, the Philippine management office, the company email and telephone details. This gives visitors a clear connection between the corporate, investor-relations and operating locations.",
    )
    action(doc, "[SCROLL TO INVESTOR RESOURCES.] ")
    say(
        doc,
        "The resource buttons use two different access rules. The Agarwood and Mango exposés remain request-gated, which supports controlled distribution. The July FAQ is public and opens directly in a new browser tab because it is the approved plain-language information resource.",
    )
    say(
        doc,
        "The draft Private Placement Memorandum and Professional Client Investment presentation are not unrestricted public downloads. Their draft, eligibility, private-placement and confidentiality conditions require controlled distribution.",
    )
    source_note(doc, "July FAQ PDF pp.11–12; Professional presentation PDF pp.33–35; draft PPM PDF pp.1–3, 16, 18–19 and 27–29; crop exposés as controlled resources.")
    transition(doc, "Before closing, I will briefly show the legal information that applies across the entire website.")

    # Legal
    script_heading(doc, "9", "Legal Pages and Footer", "Footer links: Disclaimer, Privacy Policy and Cookie Policy", "1–2 minutes")
    action(doc, "[SCROLL TO THE FOOTER. CLICK DISCLAIMER.] ")
    say(
        doc,
        "The Risk Warning and Disclaimer explains that the website is informational and is not an offer, solicitation or advice. It describes the economics as illustrative, the fund structure as proposed, the interests as long-term and illiquid, and the final transaction documents as controlling.",
    )
    say(
        doc,
        "It also summarises the principal biological, market, operational, regulatory and currency risks, together with the professional-investor and jurisdictional limitations.",
    )
    action(doc, "[RETURN TO FOOTER. OPEN PRIVACY POLICY, THEN COOKIE POLICY BRIEFLY.] ")
    say(
        doc,
        "The Privacy Policy explains what personal data is collected through enquiries, why it is processed, which systems and advisers may handle it, and what rights a visitor has. The Cookie Policy explains essential, analytics and functional cookies, including Google Analytics and Microsoft Clarity.",
    )
    say(
        doc,
        "Together, these pages make the website’s investment boundaries, data practices and tracking practices accessible from every page.",
    )
    source_note(doc, "Company Overview PDF p.2; Agarwood Exposé PDF p.23; Mango Exposé PDF p.21; draft PPM PDF pp.1 and 18–21. Privacy and cookie detail is based on website legal/operational requirements.")

    # Closing
    doc.add_page_break()
    script_heading(doc, "10", "Closing Statement", "Return to Home or Contact", "1 minute")
    action(doc, "[RETURN TO HOME HERO OR CONTACT PAGE. PAUSE.] ")
    say(
        doc,
        "To summarise, the website is structured around a clear professional-investor journey. Home introduces the fund-share proposition. Investment explains the two crop pathways and key questions. Precision Farming demonstrates the operating model. Asset Management provides optional first-hand access. Golden Forests Group establishes organisational credibility. Contact moves qualified interest into a controlled information process.",
    )
    say(
        doc,
        "Across every page, the website uses careful language because the structure, economics and final offering terms remain subject to definitive documentation. Our purpose is to provide enough information for an investor to understand the proposition and begin proper due diligence—without presenting assumptions as guarantees.",
    )
    say(doc, "Thank you. I am happy to answer any questions or return to any page in more detail.")

    # Q&A
    doc.add_heading("Presenter Q&A — Keep Available After the Walkthrough", level=1)
    qa(doc, "Do investors own the trees?", "No. Investors would own shares in the relevant fund compartment. Tree-equivalent references are for economic allocation, accounting and reporting only.")
    qa(doc, "Is USD 100,000 the final minimum?", "It is the current indicative reference subscription. The actual minimum and eligibility requirements depend on the relevant platform and definitive offering terms.")
    qa(doc, "Which mango share price is correct?", "The July documents conflict: one shows USD 371.64 and others show USD 437.08. The public website therefore does not publish a mango share price pending approved final terms.")
    qa(doc, "What returns are expected?", "The documents contain illustrative scenarios, not guarantees. Discuss harvest timing and fee mechanics, then refer the investor to approved current materials and definitive documents.")
    qa(doc, "Are the trees insured?", "The operating model expects insurance for specified risks, subject to availability, policy limits, exclusions and claims assessment. Insurance does not eliminate risk or guarantee recovery.")
    qa(doc, "What does the 20% buffer mean?", "Approximately 20% surplus replacement stock is intended to support mortality management within fund-managed inventory. Replacement timing and suitability depend on actual conditions.")
    qa(doc, "Can investors visit?", "Eligible professional and corporate investors may request a visit. Access, itinerary, timing and hospitality are confirmed individually and remain subject to availability and operating conditions.")
    qa(doc, "Why is the PPM not on the website?", "It is a draft private-placement document for controlled evaluation. It is not a final offering document and its distribution is subject to eligibility, jurisdiction and confidentiality requirements.")
    qa(doc, "Is the Q4 2026 deployment guaranteed?", "No. It is the current illustrative launch profile and remains subject to legal, administrative and operational completion.")

    doc.add_heading("Source Reference — For Presenter Preparation", level=1)
    sources = [
        "6. GF Company Overview-July 2026.pdf — PDF pp.1–2",
        "6. GF Agarwood Exposé July 2026.pdf — PDF pp.1–23",
        "6. GF Mango Exposé July 2026.pdf — PDF pp.1–21",
        "6. GF Professional Client Investment presentation - July 2026.pdf — PDF pp.2–35",
        "6. GF-FAQ-July 2026.pdf — PDF pp.1–13",
        "9. Private Placement Memorandum-Professional Investors-Draft-July 2026.pdf — PDF pp.1–31",
    ]
    for item in sources:
        doc.add_paragraph(item, style="List Bullet")
    label_box(
        doc,
        "FINAL REMINDER",
        "If a question requires a precise price, allocation, return, legal right, distribution or jurisdictional answer that is not in this script, do not improvise. Say that the point is subject to approved current materials and definitive documents, and arrange a follow-up with the appropriate team.",
    )

    ending = doc.add_paragraph()
    ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ending.paragraph_format.space_before = Pt(20)
    run = ending.add_run("End of presentation script")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(OLIVE)

    properties = doc.core_properties
    properties.title = "Golden Forests Website Presentation Script"
    properties.subject = "Presenter-ready page-by-page walkthrough of the July 2026 website"
    properties.author = "Golden Forests / OpenAI Codex"
    properties.keywords = "Golden Forests, website, presenter script, July 2026"
    properties.comments = "Main paragraphs are intended to be spoken aloud; source and Q&A boxes are presenter notes."
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(path)
    print(path.stat().st_size)
